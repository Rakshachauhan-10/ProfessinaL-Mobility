# # import fitz  # PyMuPDF
# # from docx import Document

# # def extract_text_from_pdf(pdf_path):
# #     pdf_document = fitz.open(pdf_path)
# #     extracted_text = ""

# #     for page_num in range(len(pdf_document)):
# #         page = pdf_document.load_page(page_num)  
# #         extracted_text += page.get_text()  
# #     return extracted_text

# # def save_text_to_word(text, word_path):
# #     doc = Document()
# #     doc.add_paragraph(text)
# #     doc.save(word_path)

# # pdf_path = "D:\Professional Mobility\Binocular Astronomy by Stephen Tonkin.pdf"
# # word_path = "D:\Professional Mobility\Binocular.docx"

# # text = extract_text_from_pdf(pdf_path)
# # save_text_to_word(text, word_path)

# # print(f"Text extracted from {pdf_path} and saved to {word_path}")

# ######################################################################################################################################

# import fitz  # PyMuPDF
# from docx import Document

# def extract_text_from_pdf(pdf_path):
#     pdf_document = fitz.open(pdf_path)
#     extracted_text = ""

#     for page_num in range(len(pdf_document)):
#         page = pdf_document.load_page(page_num)
#         extracted_text += page.get_text()
#     return extracted_text

# def save_text_to_word(text, word_path):
#     doc = Document()
#     # Split text into smaller chunks if it's too large
#     chunk_size = 10000  # Adjust size if necessary
#     for i in range(0, len(text), chunk_size):
#         chunk = text[i:i+chunk_size]
#         doc.add_paragraph(chunk)
#     doc.save(word_path)

# pdf_path = "D:\\Professional Mobility\\Binocular Astronomy by Stephen Tonkin.pdf"
# word_path = "D:\\Professional Mobility\\Binocular.docx"

# text = extract_text_from_pdf(pdf_path)
# save_text_to_word(text, word_path)

# print(f"Text extracted from {pdf_path} and saved to {word_path}")




import fitz  # PyMuPDF
from docx import Document

def extract_text_from_pdf(pdf_path):
    try:
        pdf_document = fitz.open(pdf_path)
        extracted_text = ""

        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            extracted_text += page.get_text()
        
        pdf_document.close()  # Close the PDF after processing
        return extracted_text
    
    except Exception as e:
        print(f"An error occurred while extracting text from the PDF: {e}")
        return None

def save_text_to_word(text, word_path):
    try:
        doc = Document()
        # Split text into smaller chunks if it's too large
        chunk_size = 10000  # Adjust size if necessary
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i+chunk_size]
            doc.add_paragraph(chunk)
        doc.save(word_path)
    
    except Exception as e:
        print(f"An error occurred while saving text to the Word document: {e}")

pdf_path = "D:\\Professional Mobility\\Agricultural pattern.pdf"
word_path = "D:\\Professional Mobility\\Agricultural pattern.docx"

text = extract_text_from_pdf(pdf_path)
if text:
    save_text_to_word(text, word_path)
    print(f"Text extracted from {pdf_path} and saved to {word_path}")



